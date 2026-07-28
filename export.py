"""
Build downloadable DOCX/PDF files from the user's EDITED text (never
straight from raw AI output -- the edit step is the safety net).

Two styles: "bullets" (tailored resume bullet points, rendered as a
bulleted list) and "letter" (cover letter, rendered as plain paragraphs
with no bullets -- a letter shouldn't look like a list).
"""
import io
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter as LETTER_SIZE
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from analysis import enforce_dash_style


def build_docx(title: str, items: list[str], style: str = "bullets") -> bytes:
    doc = Document()
    doc.add_heading(title or "Document", level=1)

    for item in items:
        clean = enforce_dash_style(item)
        if style == "letter":
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(10)
        else:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean)
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_pdf(title: str, items: list[str], style: str = "bullets") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER_SIZE,
                             leftMargin=0.8 * inch, rightMargin=0.8 * inch,
                             topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    styles = getSampleStyleSheet()
    story = [Paragraph(title or "Document", styles["Heading1"]),
             Spacer(1, 12)]

    for item in items:
        clean = enforce_dash_style(item)
        text = clean if style == "letter" else f"- {clean}"
        story.append(Paragraph(text, styles["Normal"]))
        story.append(Spacer(1, 10 if style == "letter" else 6))

    doc.build(story)
    buf.seek(0)
    return buf.read()

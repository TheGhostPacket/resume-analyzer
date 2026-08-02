"""
Renders the structured CV dict (from ai.generate_full_cv) into a clean,
professionally styled document -- a real CV layout, not just a flat
bullet list. Kept separate from export.py since the input shape and
template are meaningfully different (structured sections vs. a flat
bullet list).
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter as LETTER_SIZE
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from analysis import enforce_dash_style

# One accent color, used consistently across both formats so the
# document reads as deliberately designed rather than default-template.
ACCENT_HEX = "3F6B5E"  # deep green -- matches the app's own accent color
ACCENT_RGB = RGBColor(0x3F, 0x6B, 0x5E)
GREY_RGB = RGBColor(0x66, 0x66, 0x66)


def _clean(text: str) -> str:
    return enforce_dash_style(text or "")


def _add_bottom_border(paragraph, color_hex: str, size: int = 6):
    """Adds a thin horizontal rule under a paragraph via raw OOXML --
    python-docx has no high-level API for paragraph borders."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_cv_docx(cv: dict) -> bytes:
    doc = Document()

    # Consistent base typography for the whole document.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    if cv.get("full_name"):
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_after = Pt(2)
        run = name_p.add_run(_clean(cv["full_name"]))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = ACCENT_RGB
        run.font.name = "Calibri"

    if cv.get("contact_line"):
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(10)
        run = contact_p.add_run(_clean(cv["contact_line"]))
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREY_RGB
        _add_bottom_border(contact_p, ACCENT_HEX, size=10)

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = ACCENT_RGB
        # Letter-spacing isn't exposed by python-docx; small caps via
        # explicit uppercase text + color/weight is the reliable option.
        _add_bottom_border(p, "DDDDDD", size=4)

    if cv.get("summary"):
        add_section_heading("Summary")
        doc.add_paragraph(_clean(cv["summary"]))

    if cv.get("experience"):
        add_section_heading("Experience")
        for entry in cv["experience"]:
            header_p = doc.add_paragraph()
            header_p.paragraph_format.space_before = Pt(8)
            header_p.paragraph_format.space_after = Pt(1)
            title_run = header_p.add_run(_clean(entry.get("title", "")))
            title_run.bold = True
            title_run.font.size = Pt(11)

            org = entry.get("organization", "")
            dates = entry.get("dates", "")
            if org or dates:
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_after = Pt(4)
                meta_run = meta_p.add_run(
                    _clean(org) + ("   |   " + _clean(dates) if dates else "")
                )
                meta_run.italic = True
                meta_run.font.size = Pt(9.5)
                meta_run.font.color.rgb = GREY_RGB

            for bullet in entry.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.add_run(_clean(bullet)).font.size = Pt(10.5)

    if cv.get("education"):
        add_section_heading("Education")
        for entry in cv["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(_clean(entry.get("degree", "")))
            run.bold = True
            run.font.size = Pt(10.5)
            inst = entry.get("institution", "")
            dates = entry.get("dates", "")
            if inst or dates:
                meta_run = p.add_run(
                    "   " + _clean(inst) + ("   |   " + _clean(dates) if dates else "")
                )
                meta_run.italic = True
                meta_run.font.color.rgb = GREY_RGB
                meta_run.font.size = Pt(9.5)

    if cv.get("skills"):
        add_section_heading("Skills")
        p = doc.add_paragraph()
        p.add_run(" \u2022 ".join(_clean(s) for s in cv["skills"])).font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_cv_pdf(cv: dict) -> bytes:
    accent = colors.HexColor(f"#{ACCENT_HEX}")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER_SIZE,
                             leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                             topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    base = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=base["Heading1"], alignment=1,
                                  fontSize=22, textColor=accent, spaceAfter=2)
    contact_style = ParagraphStyle("Contact", parent=base["Normal"], alignment=1,
                                    fontSize=9.5, textColor=colors.grey, spaceAfter=10)
    section_style = ParagraphStyle("Section", parent=base["Heading2"], fontSize=12,
                                     textColor=accent, spaceBefore=16, spaceAfter=5,
                                     borderWidth=0)
    entry_title_style = ParagraphStyle("EntryTitle", parent=base["Normal"], fontSize=11,
                                         spaceBefore=8, fontName="Helvetica-Bold")
    entry_meta_style = ParagraphStyle("EntryMeta", parent=base["Normal"], fontSize=9.5,
                                        textColor=colors.grey, spaceAfter=4,
                                        fontName="Helvetica-Oblique")
    bullet_style = ParagraphStyle("Bullet", parent=base["Normal"], fontSize=10.5,
                                    leftIndent=14, spaceAfter=3)

    story = []

    if cv.get("full_name"):
        story.append(Paragraph(_clean(cv["full_name"]), name_style))
    if cv.get("contact_line"):
        story.append(Paragraph(_clean(cv["contact_line"]), contact_style))
    story.append(HRFlowable(width="100%", color=accent, thickness=1.2))

    if cv.get("summary"):
        story.append(Paragraph("SUMMARY", section_style))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#DDDDDD"), thickness=0.5))
        story.append(Spacer(1, 4))
        story.append(Paragraph(_clean(cv["summary"]), base["Normal"]))

    if cv.get("experience"):
        story.append(Paragraph("EXPERIENCE", section_style))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#DDDDDD"), thickness=0.5))
        for entry in cv["experience"]:
            story.append(Paragraph(f"<b>{_clean(entry.get('title', ''))}</b>", entry_title_style))
            org = _clean(entry.get("organization", ""))
            dates = _clean(entry.get("dates", ""))
            if org or dates:
                meta = org + ("   |   " + dates if dates else "")
                story.append(Paragraph(f"<i>{meta}</i>", entry_meta_style))
            for bullet in entry.get("bullets", []):
                story.append(Paragraph(f"- {_clean(bullet)}", bullet_style))

    if cv.get("education"):
        story.append(Paragraph("EDUCATION", section_style))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#DDDDDD"), thickness=0.5))
        for entry in cv["education"]:
            inst = _clean(entry.get("institution", ""))
            dates = _clean(entry.get("dates", ""))
            header = f"<b>{_clean(entry.get('degree', ''))}</b>"
            if inst or dates:
                header += f"   {inst}" + (f"   |   {dates}" if dates else "")
            story.append(Paragraph(header, entry_title_style))

    if cv.get("skills"):
        story.append(Paragraph("SKILLS", section_style))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#DDDDDD"), thickness=0.5))
        story.append(Spacer(1, 4))
        story.append(Paragraph(" &bull; ".join(_clean(s) for s in cv["skills"]), base["Normal"]))

    doc.build(story)
    buf.seek(0)
    return buf.read()

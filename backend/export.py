"""
Build downloadable DOCX/PDF files from the user's EDITED bullet text
(never straight from raw AI output -- the edit step is the safety net).
"""
import io
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from analysis import enforce_dash_style


def build_docx(title: str, bullets: list[str]) -> bytes:
    doc = Document()
    heading = doc.add_heading(title or "Tailored Resume Section", level=1)
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(enforce_dash_style(bullet))
        run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_pdf(title: str, bullets: list[str]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                             leftMargin=0.8 * inch, rightMargin=0.8 * inch,
                             topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    styles = getSampleStyleSheet()
    story = [Paragraph(title or "Tailored Resume Section", styles["Heading1"]),
             Spacer(1, 12)]

    for bullet in bullets:
        clean = enforce_dash_style(bullet)
        story.append(Paragraph(f"- {clean}", styles["Normal"]))
        story.append(Spacer(1, 6))

    doc.build(story)
    buf.seek(0)
    return buf.read()

"""
ATS formatting checker.

Distinct from keyword matching entirely -- this looks at the STRUCTURE
of the uploaded file, not its content. A resume can have a perfect
keyword match and still get auto-rejected because a real ATS parser
misreads its layout. Each check below is a real, well-documented ATS
failure mode, not a guess -- and each finding says exactly what was
detected and why it matters, so nothing here is a black-box score.
"""
import io


def check_docx_ats(file_bytes: bytes) -> list[dict]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    issues = []

    if len(doc.tables) > 0:
        issues.append({
            "issue": f"Contains {len(doc.tables)} table(s)",
            "why": "Many ATS parsers read table content out of order or "
                   "skip it entirely. Consider using plain paragraphs "
                   "instead of tables for layout.",
        })

    if len(doc.inline_shapes) > 0:
        issues.append({
            "issue": f"Contains {len(doc.inline_shapes)} embedded image(s)/graphic(s)",
            "why": "Text inside images (like a graphic skills chart or a "
                   "logo with embedded text) is invisible to ATS parsers "
                   "entirely -- it will not be read at all.",
        })

    # Check for multi-column section layout via the section's raw XML
    # (python-docx doesn't expose column count as a plain property).
    for section in doc.sections:
        sect_pr = section._sectPr
        cols = sect_pr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
        )
        if cols is not None:
            num_attr = cols.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
            )
            if num_attr and int(num_attr) > 1:
                issues.append({
                    "issue": f"Multi-column layout detected ({num_attr} columns)",
                    "why": "Many ATS systems read multi-column resumes "
                           "left-to-right across the whole page rather than "
                           "column-by-column, scrambling the reading order.",
                })
                break

    header_footer_text = ""
    for section in doc.sections:
        header_footer_text += "".join(p.text for p in section.header.paragraphs)
        header_footer_text += "".join(p.text for p in section.footer.paragraphs)
    if header_footer_text.strip():
        issues.append({
            "issue": "Header/footer contains text",
            "why": "Some ATS systems ignore header/footer content entirely "
                   "-- if key info (like contact details) only lives there, "
                   "move it into the main document body.",
        })

    return issues


def check_pdf_ats(file_bytes: bytes) -> list[dict]:
    import pdfplumber

    issues = []
    table_count = 0
    image_count = 0
    multi_column_pages = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            tables = page.find_tables()
            if tables:
                table_count += len(tables)

            if page.images:
                image_count += len(page.images)

            # Heuristic multi-column detection: cluster word x0 positions.
            # If words consistently fall into two distinct, separated
            # horizontal bands across many lines, that's a strong signal
            # of a two-column layout (very common ATS failure mode).
            words = page.extract_words()
            if len(words) > 20:
                page_width = page.width
                left_band = sum(1 for w in words if w["x0"] < page_width * 0.45)
                right_band = sum(1 for w in words if w["x0"] > page_width * 0.55)
                if left_band > 10 and right_band > 10:
                    multi_column_pages += 1

    if table_count > 0:
        issues.append({
            "issue": f"Contains {table_count} table(s)",
            "why": "Many ATS parsers read table content out of order or "
                   "skip it entirely. Consider using plain paragraphs "
                   "instead of tables for layout.",
        })

    if image_count > 0:
        issues.append({
            "issue": f"Contains {image_count} embedded image(s)/graphic(s)",
            "why": "Text inside images (like a graphic skills chart or a "
                   "logo with embedded text) is invisible to ATS parsers "
                   "entirely -- it will not be read at all.",
        })

    if multi_column_pages > 0:
        issues.append({
            "issue": f"Likely multi-column layout detected on {multi_column_pages} page(s)",
            "why": "Many ATS systems read multi-column resumes "
                   "left-to-right across the whole page rather than "
                   "column-by-column, scrambling the reading order. "
                   "(Heuristic detection -- worth a manual double-check.)",
        })

    return issues


def check_ats_formatting(filename: str, file_bytes: bytes) -> dict:
    lower = filename.lower()
    if lower.endswith(".docx"):
        issues = check_docx_ats(file_bytes)
    elif lower.endswith(".pdf"):
        issues = check_pdf_ats(file_bytes)
    else:
        issues = []

    return {
        "issues": issues,
        "clean": len(issues) == 0,
    }
